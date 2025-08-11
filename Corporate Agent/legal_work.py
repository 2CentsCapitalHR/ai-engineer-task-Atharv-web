from google import genai
from docx import Document
import os,json
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
import shutil
from dotenv import load_dotenv
load_dotenv()

# Configure Gemini API
client = genai.Client()
MODEL = "gemini-2.5-flash"
docx_file_path = "input" # folder that consists of only docx files.

def identify_legal_process(docx_file_paths):
    # Read and concatenate document text
    all_text = []
    for file_path in docx_file_paths:
        doc = Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        all_text.append(f"--- {os.path.basename(file_path)} ---\n{text}")
    
    combined_text = "\n\n".join(all_text)
    
    # Create prompt for LLM
    prompt = f"""
    You are an ADGM legal assistant.
    Based on the following document contents, identify the most likely legal process that the user is attempting under ADGM jurisdiction.

    respond with only the identified legal process.
    Documents:
    {combined_text[:2000]}
    """
    
    # Call Gemini API
    response = client.models.generate_content(
        model=MODEL,
        contents=prompt,
    )
    
    return response.text.strip()


def compare_with_checklist(legal_process, uploaded_docx_paths, checklist_path):
    # Load checklist.json
    with open(checklist_path, 'r', encoding='utf-8') as f:
        checklist = json.load(f)

    # Find the required document list for this legal process
    required_docs = []
    for category in ["companies", "partnerships", "cell_and_investment_companies","employment"]:
        if category in checklist and legal_process in checklist[category]:
            for doc_entry in checklist[category][legal_process]["documents"]:
                if isinstance(doc_entry, dict) and "text" in doc_entry:
                    required_docs.append(doc_entry["text"])
                elif isinstance(doc_entry, str):
                    required_docs.append(doc_entry)

    # Extract uploaded file names
    uploaded_file_names = [os.path.basename(p) for p in uploaded_docx_paths]
    missing_docs = []

    for req_doc in required_docs:
        match_found = False
        
        # Check each uploaded file with LLM
        for fname in uploaded_file_names:
            prompt = f"""
            You are a legal assistant.
            Given the required ADGM document name and an uploaded file name, 
            determine if they represent the same type of document, even if the names differ.
            Answer ONLY with 'YES' or 'NO'.

            Required document: {req_doc}
            Uploaded file name: {fname}
            """
            
            try:
                response = genai.generate_content(
                    model=MODEL,
                    contents=prompt,
                    generation_config={"temperature": 0.0, "max_output_tokens": 5}
                )
                answer = response.text.strip().upper()
                if answer == "YES":
                    match_found = True
                    break
            except Exception as e:
                print(f"LLM match check failed for {fname}: {e}")
        
        if not match_found:
            missing_docs.append(req_doc)

    return {
        "num_docs_uploaded": len(uploaded_docx_paths),
        "num_docs_required": len(required_docs),
        "missing_docs": missing_docs,
        "required_docs": required_docs,
    }


db_path = "faiss_db"

def load_faiss_db(db_path):
    embeddings = OllamaEmbeddings(model = 'all-minilm:22m')
    return FAISS.load_local(db_path, embeddings, allow_dangerous_deserialization=True)

def extract_docx_chunks(file_path):
    doc = Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ". ", " "]
    )

    chunks = splitter.split_text(full_text)
    return chunks

def check_legal_issues(docx_paths, faiss_db_path):

    json_array = {
        "section": "Short summary for the chunk",
        "issue": "Short description of the problem",
        "severity": "Low / Medium / High",
        "suggestion": "Clear legal recommendation to fix"
    }

    vectorstore = load_faiss_db(faiss_db_path)
    retriever = vectorstore.as_retriever(search_kwargs={"k": 5})

    prompt_template = """
    You are an ADGM legal compliance assistant.
    You will receive:
    1. A chunk of text from a user-uploaded legal document.
    2. Relevant ADGM legal context retrieved from official regulations.

    Task:
    - Identify any legal red flags or inconsistencies in this chunk.
    - If compliant, return Nothing Wrong.
    - If issues exist, respond in valid JSON format: {json_format}

    Relevant ADGM legal context:
    {context}

    """

    prompt = PromptTemplate(template=prompt_template, input_variables=["chunk_text", "context","json_format"])

    results = []

    for file_path in docx_paths:
        chunks = extract_docx_chunks(file_path)
        for chunk in chunks:
            retrieved_docs = retriever.invoke(chunk)
            context = "\n".join([doc.page_content for doc in retrieved_docs])

            final_prompt = prompt.format(json_format=json_array, context=context)
            response = client.models.generate_content(
                    model=MODEL,
                    contents=final_prompt,
            )
            
            raw_text = response.text
            cleaned_text = raw_text.strip().lower() 
            if cleaned_text == "nothing wrong":
                return []

            else:
                try:
                    issues = json.loads(response.text)
                    for issue in issues:
                        results.append({
                            "document": os.path.basename(file_path),
                            "section": issue.get("section"),
                            "issue": issue.get("issue"),
                            "severity": issue.get("severity"),
                            "suggestion": issue.get("suggestion")
                        })
                except json.JSONDecodeError:
                    print(f"LLM returned invalid JSON for chunk in {file_path}")
    return results if results else "No issues"

def annotate_and_generate_report(
    process, # process identified by LLM
    uploaded_docx_paths, # input
    checklist_result, 
    checklist_path, # checklist.json
    issues_found, # from function check_legal_issues
    output_folder # output
):

    os.makedirs(output_folder, exist_ok=True)
    uploaded_file_names = [os.path.basename(p) for p in uploaded_docx_paths]

    if issues_found == "No issues":
            report = {
            "process": process,
            "documents_uploaded": uploaded_file_names,
            "required_documents": checklist_result['required_docs'],
            "missing_documents": checklist_result['missing_docs'],
            "issues_found": issues_found
    }
        
    else:
        # Load checklist
        with open(checklist_path, "r", encoding="utf-8") as f:
            checklist = json.load(f)

        # Get required documents for process
        required_docs = []
        for category in checklist:
            if process in checklist[category]:
                for doc_entry in checklist[category][process]["documents"]:
                    if isinstance(doc_entry, dict) and "text" in doc_entry:
                        required_docs.append(doc_entry["text"])
                    elif isinstance(doc_entry, str):
                        required_docs.append(doc_entry)

        missing_docs = [doc for doc in required_docs if doc not in uploaded_file_names]

        # Group issues by document
        issues_by_doc = {}

        for issue in issues_found:
            doc_name = issue.get("document")
            if doc_name not in issues_by_doc:
                issues_by_doc[doc_name] = []
            issues_by_doc[doc_name].append(issue)

        # Annotate each file
        for file_path in uploaded_docx_paths:
            doc_name = os.path.basename(file_path)
            related_issues = issues_by_doc.get(doc_name, [])

            doc = Document(file_path)

            if related_issues:
                # Highlight flagged sections
                for para in doc.paragraphs:
                    for issue in related_issues:
                        if issue["section"] and issue["section"].lower() in para.text.lower():
                            note = f"[ISSUE: {issue['issue']}] [SEVERITY: {issue['severity']}] [SUGGESTION: {issue['suggestion']}]"
                            run = para.add_run(f"\n{note}")
                            run.font.highlight_color = 6  # Yellow
            # Save updated document to output
            doc.save(os.path.join(output_folder, doc_name))
        
        report = {
            "process": process,
            "documents_uploaded": uploaded_file_names,
            "required_documents": required_docs,
            "missing_documents": missing_docs,
            "issues_found": issues_found,
        }

    # Copy unchanged files (if no issues)
    for file_path in uploaded_docx_paths:
        out_path = os.path.join(output_folder, os.path.basename(file_path))
        if not os.path.exists(out_path):  # means it wasn't saved in annotation loop
            shutil.copy(file_path, out_path)

    return report